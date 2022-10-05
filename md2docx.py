from typing import List, Dict, Final
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from sympy import preview
from tqdm import tqdm
import yaml
import os
import os.path
import re
import json
import regex
import subprocess
import sys

filepath: str = ""
destpath: str = ""

if len(sys.argv) == 1:
    print("source markdown file is required")
    exit(1)
elif len(sys.argv) == 2:
    filepath = sys.argv[1]
    destpath = sys.argv[1].replace(".md", ".docx")
else:
    filepath = sys.argv[1]
    destpath = sys.argv[2]

filepath = os.path.join(os.getcwd(), filepath)
destpath = os.path.join(os.getcwd(), destpath)
configpath = os.path.join(os.path.dirname(sys.argv[0]), "./md2docx-config.yml")

with open(filepath, "rb") as f:
    markdown = f.read().decode('utf-8')
with open(configpath, "r", encoding="utf-8") as f:
    global_config = yaml.safe_load(f)
print(global_config)
lines = markdown.splitlines()
i = 0

# YAMLヘッダの解析
read_yaml_header = False
yaml_header_lines: List[str] = []
while i < len(lines) and (read_yaml_header == True or
                          len(yaml_header_lines) == 0):

    if lines[i] == "":
        pass
    elif lines[i] == "---":
        read_yaml_header = (not read_yaml_header)
    elif read_yaml_header == True:
        yaml_header_lines.append(lines[i])

    i += 1
config: Final = yaml.safe_load("\n".join(yaml_header_lines))

LINE_TYPE_EMPTY = "empty"
LINE_TYPE_NEWPAGE = "newpage"
LINE_TYPE_LATEX = "latex"
LINE_TYPE_MERMAID = "mermaid"
LINE_TYPE_IMAGE = "image"
LINE_TYPE_LIST = "list"
LINE_TYPE_HEADING = "heading"
LINE_TYPE_REFERENCE = "reference"
LINE_TYPE_PARAGRAPH = "paragraph"


class LineInfo:

    def __init__(self, kind: str, content: Dict[str, str]):
        self.kind = kind
        self.content = content


line_infos: List[LineInfo] = []
# Markdownの解析
with tqdm(total=len(lines), desc="parsing markdown") as progress:
    while i < len(lines):
        progress.update(i)
        if lines[i] == "":
            line_infos.append(LineInfo(LINE_TYPE_EMPTY, {}))
        elif lines[i] == "---":
            line_infos.append(LineInfo(LINE_TYPE_NEWPAGE, {}))
        elif lines[i] == "$$":
            content: str = ""
            i += 1
            while i < len(lines) and lines[i] != "$$":
                content += lines[i]
                i += 1
            content = f"$$\n{content}\n$$"
            line_infos.append(LineInfo(LINE_TYPE_LATEX, {"content": content}))
        elif lines[i].startswith("```"):
            if lines[i][3:].startswith("mermaid"):
                matches = re.match("\[([^\[\]]+)\]", lines[i][10:])
                if matches is not None:
                    content = ""
                    i += 1
                    while i < len(lines) and lines[i] != "```":
                        content += lines[i] + "\n"
                        i += 1
                    line_infos.append(
                        LineInfo(LINE_TYPE_MERMAID, {
                            "content": content,
                            "desc": matches.groups()[0],
                        }))

        elif re.fullmatch("(\*|-|\+) \s+", lines[i]) is not None:
            line_infos.append(
                LineInfo(LINE_TYPE_LIST, {
                    "star": lines[i][:1],
                    "content": lines[i][2:]
                }))
        elif re.fullmatch("#+ .*", lines[i]) is not None:
            level = lines[i].find(" ")
            line_infos.append(
                LineInfo(LINE_TYPE_HEADING, {
                    "level": str(level),
                    "content": lines[i][level + 1:]
                }))
        elif re.fullmatch("\[\^[^\[\]]+\]: .*", lines[i]) is not None:
            key_end_idx = lines[i].find("]: ")
            line_infos.append(
                LineInfo(
                    LINE_TYPE_REFERENCE, {
                        "key": lines[i][1:key_end_idx],
                        "content": lines[i][key_end_idx + 3:]
                    }))
        elif re.fullmatch("!\[.+\]\([^\(\)]+\)", lines[i]) is not None:
            desc_end_idx = lines[i].find("](")
            line_infos.append(
                LineInfo(
                    LINE_TYPE_IMAGE, {
                        "desc": lines[i][2:desc_end_idx],
                        "path": lines[i][desc_end_idx + 2:len(lines[i]) - 1]
                    }))
        else:
            line_infos.append(
                LineInfo(LINE_TYPE_PARAGRAPH, {"content": lines[i][0:]}))

        i += 1

# 中間生成物の解析とDocxへの出力
jpn = "[\p{Hiragana}\p{Katakana}\p{Han}ー（）．，]+"


def split_jpn(expr: str) -> List[Dict[str, str]]:
    founds = regex.findall(jpn, expr)
    i = 0
    results: List[Dict[str, str]] = []
    for found in founds:
        jpn_idx = expr.find(found, i)
        if jpn_idx != i:
            results.append({"content": expr[i:jpn_idx], "region": "europe"})
        results.append({
            "content": expr[jpn_idx:jpn_idx + len(found)],
            "region": "jpn"
        })
        i = jpn_idx + len(found)
    if i < len(expr):
        results.append({
            "content": expr[i:],
            "region": "europe",
        })

    return results


replace_refered_keydict: Dict[str, str] = {}
replace_refered_valuedict: Dict[str, str] = {}


def replace_expr(expr: str, para: Paragraph, is_heading: bool = False):
    global replace_refered_keydict
    global replace_refered_valuedict

    prev_end = 0
    for found in re.finditer("\[\^([^\[\]]+)\]", expr):
        if prev_end != found.start():
            spliteds = split_jpn(expr[prev_end:found.start()])
            for splited in spliteds:
                para.add_run(splited["content"]).font.name = select_font(
                    splited["region"], is_heading=is_heading)

        value = replace_refered_keydict.get(found.group(0))
        if value is None:
            continue
        index = replace_refered_valuedict[value]
        if index is None:
            continue
        spliteds = split_jpn(f"[{index}]")
        for splited in spliteds:
            run = para.add_run(splited["content"])
            run.font.name = select_font(splited["region"],
                                        is_heading=is_heading)
            run.font.superscript = True
        prev_end = found.end()

    spliteds = split_jpn(expr[prev_end:])
    for splited in spliteds:
        para.add_run(splited["content"]).font.name = select_font(
            splited["region"], is_heading=is_heading)
    return expr


def select_font(region: str, is_heading: bool = False) -> str:
    fonts = {
        "jpn": {
            True: "ＭＳ ゴシック",
            False: "ＭＳ 明朝",
        },
        "europe": {
            True: "Arial",
            False: "Times New Roman",
        }
    }
    return fonts[region][is_heading]


i = 0
i_image = 1
i_reference = 1
i_headings = [0] * 9
dest = Document()
refereds: List[str] = []

with tqdm(total=len(line_infos), desc="preparing to generate docx") as progress:

    paragraph = dest.add_paragraph("", "Title")
    replace_expr(config["Title"], paragraph, is_heading=True)
    paragraph = dest.add_paragraph("")
    spliteds = split_jpn(replace_expr(config["Author"], paragraph))

    while i < len(line_infos):

        progress.total = len(line_infos)
        progress.update(i)

        if line_infos[i].kind == LINE_TYPE_IMAGE:
            description = f"Figure {i_headings[0]}.{i_image}. {line_infos[i].content['desc']}"
            path = os.path.join(os.path.dirname(filepath),
                                line_infos[i].content["path"])

            paragraph = dest.add_paragraph("")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.add_run().add_picture(path, width=Inches(5))
            paragraph = dest.add_paragraph("")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            replace_expr(description, paragraph)
            i_image += 1

        if line_infos[i].kind == LINE_TYPE_EMPTY:
            pass

        if line_infos[i].kind == LINE_TYPE_NEWPAGE:
            dest.add_page_break()

        if line_infos[i].kind == LINE_TYPE_PARAGRAPH:
            paragraph = dest.add_paragraph("")
            replace_expr(line_infos[i].content["content"], paragraph)

        if line_infos[i].kind == LINE_TYPE_HEADING:

            level = int(line_infos[i].content["level"])
            i_headings[level - 1] += 1
            for i_heading in range(level, len(i_headings)):
                i_headings[i_heading] = 0

            paragraph = dest.add_paragraph("", f"""Heading {level}""")

            numbering = f"{i_headings[0]}"
            for i_heading in i_headings[1:level]:
                numbering = f"{numbering}.{i_heading}"
            replace_expr(f"""{numbering}. {line_infos[i].content["content"]}""",
                         paragraph,
                         is_heading=True)

        if line_infos[i].kind == LINE_TYPE_LATEX:
            content = line_infos[i].content["content"]
            image_bin = BytesIO()
            preview(content,
                    viewer="BytesIO",
                    eular=False,
                    fontsize=11,
                    outputbuffer=image_bin)

            paragraph = dest.add_paragraph("")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.add_run().add_picture(image_bin)

        if line_infos[i].kind == LINE_TYPE_MERMAID:
            content = line_infos[i].content["content"]
            mmd_filename = f"./{i}.mmd"
            filename = f"./{i}.png"
            with open(mmd_filename, "xt", encoding="utf-8") as f:
                print(content, file=f)
            with open("puppeteer-config.json", "xt", encoding="utf-8") as f:
                print(json.dumps({
                    "executablePath": global_config["chromePath"],
                    "args": global_config["puppeteerArgs"],
                }),
                      file=f)
            with open("mermaid.css", "xt", encoding="utf-8") as f:
                print("""
                .mermaid {
                    height: 60% !important;
                }
                .label {
                    font-size: 10.5px !important;
                }
                """,
                      file=f)
            completed = subprocess.run([
                "mmdc",
                "-i",
                mmd_filename,
                "-o",
                filename,
                "-p",
                "puppeteer-config.json",
                "-C",
                "mermaid.css",
                "-H",
                "1200",
            ])
            completed.check_returncode()
            description = f"Figure {i_headings[0]}.{i_image}. {line_infos[i].content['desc']}"

            paragraph = dest.add_paragraph("")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.add_run().add_picture(filename)
            paragraph = dest.add_paragraph("")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            replace_expr(description, paragraph)
            i_image += 1

            os.remove(mmd_filename)
            os.remove(filename)
            os.remove("puppeteer-config.json")
            os.remove("mermaid.css")

        if line_infos[i].kind == LINE_TYPE_REFERENCE:
            key = f"[{line_infos[i].content['key']}]"
            value = line_infos[i].content["content"]
            replace_refered_keydict[key] = value
            if replace_refered_valuedict.get(value) == None:
                replace_refered_valuedict[value] = str(
                    len(replace_refered_valuedict) + 1)
                refereds.append(value)

        i += 1

        if i >= len(line_infos) and len(refereds) > 0:
            line_infos.append(
                LineInfo(LINE_TYPE_HEADING, {
                    "level": "1",
                    "content": "参考文献",
                }))
            for idx, refered in enumerate(refereds):
                line_infos.append(
                    LineInfo(LINE_TYPE_PARAGRAPH,
                             {"content": f"[{idx+1}]: {refered}"}))
            refereds = []

# Docxへの出力
print(f"saving docx file ({destpath})...")
dest.save(destpath)
